from __future__ import annotations

import hashlib
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "รายงานโครงงานเกม 2D.docx"
OUTPUT = ROOT / "รายงานโครงงานเกม 2D_เสพติดการโดดเรียน_Final.docx"
EXPECTED_SHA256 = "9c85f211db4abba21a9a1ac0d4dbb6db3b29d2f5ea8dbcec105130e1f4eda7f3"
GAME_URL = "https://avocadough.github.io/skip-class-addicted/"
REPO_URL = "https://github.com/Avocadough/skip-class-addicted"
RELEASE_URL = "https://github.com/Avocadough/skip-class-addicted/releases/tag/v0.1.0-demo"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_run_font(run, size: float, bold=False, color="18233B") -> None:
    run.font.name = "Sarabun"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Sarabun")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Sarabun")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Sarabun")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, space_after=5, line_spacing=1.08) -> None:
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_text(document, text: str, size=11.5, bold=False, color="18233B", align=None, after=5):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    style_paragraph(paragraph, after)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, color)
    return paragraph


def add_heading(document, number: str, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"{number}  {title}")
    set_run_font(run, 16, True, "246B8E")
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "F3C969")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.24)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    style_paragraph(paragraph, 3)
    run = paragraph.add_run(f"•  {text}")
    set_run_font(run, 11.2)


def add_image(document, path: Path, caption: str, width=Inches(6.35)) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=width)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cap, 7)
    run = cap.add_run(caption)
    set_run_font(run, 9.5, False, "53647D")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "246B8E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    run.append(run_props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_qr(url: str, output: Path) -> None:
    qr = qrcode.QRCode(version=None, box_size=8, border=2)
    qr.add_data(url)
    qr.make(fit=True)
    qr.make_image(fill_color="#18233B", back_color="white").save(output)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def main() -> None:
    digest = hashlib.sha256(TEMPLATE.read_bytes()).hexdigest()
    if digest != EXPECTED_SHA256:
        raise SystemExit(f"Template hash mismatch: {digest}")

    qa = ROOT / "qa-output" / "report-assets"
    qa.mkdir(parents=True, exist_ok=True)
    gameplay_qr = qa / "gameplay-qr.png"
    repo_qr = qa / "repository-qr.png"
    create_qr(GAME_URL, gameplay_qr)
    create_qr(REPO_URL, repo_qr)

    document = Document(TEMPLATE)
    clear_body(document)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    title = add_text(document, "รายงานโครงงานเกม 2D", 25, True, "18233B", WD_ALIGN_PARAGRAPH.CENTER, 3)
    title.paragraph_format.space_before = Pt(4)
    add_text(document, "เสพติดการโดดเรียน – Skip class addicted", 21, True, "246B8E", WD_ALIGN_PARAGRAPH.CENTER, 2)
    add_text(document, "2D Stealth-Puzzle Platformer Demo", 12.5, False, "53647D", WD_ALIGN_PARAGRAPH.CENTER, 10)

    meta = document.add_table(rows=4, cols=2)
    meta.autofit = False
    meta.columns[0].width = Inches(1.65)
    meta.columns[1].width = Inches(5.15)
    values = [
        ("รายวิชา", "CP410844 Computer Game Development"),
        ("ภาคการศึกษา", "ภาคการศึกษาต้น ปีการศึกษา 2569"),
        ("กลุ่ม", "กลุ่มที่ 19 • สาขาวิทยาการคอมพิวเตอร์"),
        ("สถานะ", "Web Demo v0.1.0 • 3 ด่าน • เวลาเล่นประมาณ 3–5 นาที"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        for cell in row.cells:
            set_cell_margins(cell)
        set_cell_shading(row.cells[0], "246B8E")
        set_cell_shading(row.cells[1], "EEF3F7")
        row.cells[0].text = ""
        row.cells[1].text = ""
        r0 = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(r0, 10.5, True, "FFFFFF")
        r1 = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r1, 10.5)

    add_text(document, "จัดทำโดย", 13, True, "18233B", None, 4)
    members = document.add_table(rows=4, cols=3)
    members.autofit = False
    headings = ["รหัสนักศึกษา", "ชื่อ–สกุล", "สาขาวิชา"]
    rows = [
        ("663380252-6", "กษิเดช สุขศีล", "วิทยาการคอมพิวเตอร์"),
        ("663380531-2", "ภูริณัฐ ศรีไตรรัตน์", "วิทยาการคอมพิวเตอร์"),
        ("663380262-3", "ณภัตร ช้อยกิ่ง", "วิทยาการคอมพิวเตอร์"),
    ]
    for col, heading in enumerate(headings):
        cell = members.rows[0].cells[col]
        set_cell_shading(cell, "18233B")
        cell.text = ""
        run = cell.paragraphs[0].add_run(heading)
        set_run_font(run, 10.2, True, "FFFFFF")
    for row_index, values_row in enumerate(rows, 1):
        for col, value in enumerate(values_row):
            cell = members.rows[row_index].cells[col]
            set_cell_shading(cell, "F5F7FA" if row_index % 2 else "EAF0F4")
            set_cell_margins(cell)
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, 10)

    document.add_page_break()
    add_heading(document, "1.", "ธีม หรือ แนวเกม (Game Genres)")
    add_text(document, "เกม 2D แบบ Side-scrolling แนว Stealth-Puzzle Platformer ผสมการกระโดด การลอบเร้น และปริศนาไอเท็ม นำเสนอผ่านบรรยากาศโรงเรียนไทยและอารมณ์ขันเชิงเสียดสีที่เหมาะกับผู้เล่นอายุ 12 ปีขึ้นไป")

    add_heading(document, "2.", "กลุ่มเป้าหมาย")
    add_bullet(document, "นักเรียน นักศึกษา และผู้เล่นทั่วไปที่ชอบเกมสั้น เล่นจบได้ในหนึ่งรอบ")
    add_bullet(document, "ผู้เล่น Desktop และ Mobile browser แนวนอน โดยไม่ต้องติดตั้งเกม")
    add_bullet(document, "ผู้เล่นที่ชอบ Stealth แบบอ่านสถานการณ์ได้ชัดเจนและ Retry ได้รวดเร็ว")

    add_heading(document, "3.", "เนื้อเรื่องย่อ")
    add_text(document, "“น็อต” นักเรียนที่อยากหนีออกจากคาบเรียน ต้องผ่านทางเดินอาคารเรียน โรงอาหาร และห้องสมุดก่อนประตูโรงเรียนถูกล็อก ระหว่างทางเขาต้องหลบครูเวร หัวหน้าห้อง บรรณารักษ์ และกล้องวงจรปิด พร้อมใช้สิ่งของรอบตัวสร้างข้ออ้างและเปิดทางหนีให้สำเร็จ")

    add_heading(document, "4.", "รูปแบบการเล่น และ กติกา")
    for text in [
        "เป้าหมาย: ไปถึง EXIT ของแต่ละด่านก่อนเวลาหมด ด่านใช้เวลา 60, 75 และ 90 วินาทีตามลำดับ",
        "ค่าความสงสัยเพิ่มเมื่ออยู่ในมุมมองครูหรือ CCTV เมื่อพ้นสายตาจะค่อย ๆ ลดลง และแพ้ทันทีเมื่อถึง 100",
        "การวิ่งและการลงพื้นสร้างเสียง ทำให้ NPC เปลี่ยนจาก Patrol เป็น Investigate หรือ Chase",
        "ผู้เล่นย่อหรือซ่อนในตู้/หลังสิ่งกำบังได้ และใช้ Inventory สามช่องสำหรับเหรียญ ขนม การบ้านปลอม และกุญแจ",
        "เมื่อแพ้จะเริ่มด่านปัจจุบันใหม่ ไม่ย้อนกลับไปเริ่มเกมทั้งหมด",
    ]:
        add_bullet(document, text)

    controls = document.add_table(rows=1, cols=2)
    controls.rows[0].cells[0].text = "Keyboard"
    controls.rows[0].cells[1].text = "การกระทำ"
    for cell in controls.rows[0].cells:
        set_cell_shading(cell, "246B8E")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 10.5, True, "FFFFFF")
    for keys, action in [
        ("A/D หรือ ←/→", "เดิน"), ("Shift", "วิ่ง – เร็วแต่มีเสียง"),
        ("Space", "กระโดด"), ("S หรือ ↓", "ย่อ/ซ่อน"),
        ("E", "โต้ตอบ"), ("Q", "ใช้ไอเท็ม"), ("1–3", "เลือกช่องไอเท็ม"),
    ]:
        cells = controls.add_row().cells
        cells[0].text, cells[1].text = keys, action
        for cell in cells:
            set_cell_margins(cell, 65, 100, 65, 100)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, 9.8)

    document.add_page_break()
    add_heading(document, "5.", "ตัวละคร")
    for text in [
        "น็อต — ตัวละครผู้เล่น นักเรียนชายที่ต้องใช้การสังเกตและไหวพริบเพื่อหนีเรียน",
        "ครูเวร / ครูฝ่ายปกครอง — เดินตรวจ เมื่อเห็นผู้เล่นจะไล่ตามและเพิ่มค่าความสงสัย",
        "หัวหน้าห้องมายด์ — ตรวจสอบเสียงและถามเหตุผลในห้องสมุด ผู้เล่นต้องมีการบ้านปลอมจึงตอบผ่าน",
        "บรรณารักษ์และครูห้องสมุด — เฝ้าเส้นทางกุญแจและทางออก",
        "CCTV — เซนเซอร์หมุนกวาด หากผู้เล่นอยู่ในลำแสงจะเพิ่มค่าความสงสัย",
    ]:
        add_bullet(document, text)

    add_heading(document, "6.", "ไอเท็ม ปริศนา และ กับดัก")
    items = document.add_table(rows=1, cols=3)
    for index, text in enumerate(["ไอเท็ม/ระบบ", "การใช้งาน", "ด่าน"]):
        items.rows[0].cells[index].text = text
        set_cell_shading(items.rows[0].cells[index], "18233B")
        for run in items.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, 10.2, True, "FFFFFF")
    for values_row in [
        ("ตู้/โต๊ะซ่อน", "กด E เพื่อตัดการตรวจจับชั่วคราว", "1–3"),
        ("เหรียญ 10 บาท", "ใช้ซื้ิอขนมจากตู้ขาย", "2"),
        ("ขนม", "กด Q เพื่อสร้างเสียงล่อ NPC", "2"),
        ("การบ้านปลอม", "ปลดคำตอบข้ออ้างที่ถูกต้อง", "3"),
        ("Storage Key", "เปิดประตู EXIT สุดท้าย", "3"),
        ("CCTV / มุมมองครู", "กับดักเพิ่มค่าความสงสัย", "1–3"),
    ]:
        cells = items.add_row().cells
        for index, value in enumerate(values_row):
            cells[index].text = value
            set_cell_margins(cells[index], 70, 90, 70, 90)
            for run in cells[index].paragraphs[0].runs:
                set_run_font(run, 9.6)

    add_heading(document, "7.", "ด่าน/ฉาก")
    add_text(document, "เดโมออกแบบให้รอบเล่นสำเร็จใช้เวลาประมาณ 3–5 นาที โดยแต่ละด่านแนะนำระบบใหม่เพียงหนึ่งกลุ่มและยังคงมีทางเลือกในการผ่าน")
    screenshots = [
        (ROOT / "qa-output" / "gameplay-level-1.png", "ภาพที่ 1 ทางเดินอาคารเรียน — ตู้ซ่อน เส้นทางกระโดด และครูเวร"),
        (ROOT / "qa-output" / "gameplay-level-2.png", "ภาพที่ 2 โรงอาหาร — ตู้ขายขนม ระบบเสียง และหัวหน้าห้อง"),
        (ROOT / "qa-output" / "gameplay-level-3.png", "ภาพที่ 3 ห้องสมุด — CCTV ชั้นหนังสือ บทสนทนา และกุญแจ"),
    ]
    for path, caption in screenshots:
        add_image(document, path, caption)

    add_heading(document, "8.", "ประโยชน์ของเกม")
    for text in [
        "ฝึกการสังเกตรูปแบบการเดิน การวางแผนเส้นทาง และการตัดสินใจภายใต้เวลา",
        "สาธิตการออกแบบระบบเกมด้วย State machine, Signal, Collision query และ Responsive UI ใน Godot",
        "เผยแพร่แบบ Web ช่วยให้ทดลองเล่นและประเมินผลงานได้ง่ายโดยไม่ต้องติดตั้งโปรแกรม",
        "ทะเบียน Asset และสัญญาอนุญาตช่วยสร้างแนวปฏิบัติที่ถูกต้องด้านลิขสิทธิ์ซอฟต์แวร์",
    ]:
        add_bullet(document, text)

    add_heading(document, "9.", "การเผยแพร่และแหล่งอ้างอิง")
    links = [("Gameplay / GitHub Pages", GAME_URL), ("Source repository", REPO_URL), ("Web ZIP release", RELEASE_URL)]
    for label, url in links:
        paragraph = document.add_paragraph()
        style_paragraph(paragraph, 5)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, 10.5, True)
        add_hyperlink(paragraph, url, url)

    qr_table = document.add_table(rows=1, cols=2)
    for cell, qr_path, caption in zip(qr_table.rows[0].cells, [gameplay_qr, repo_qr], ["เล่นเกม (Gameplay)", "Repository"]):
        set_cell_margins(cell, 80, 80, 80, 80)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(qr_path), width=Inches(1.35))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        set_run_font(run, 10, True, "246B8E")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    footer_run = footer.add_run("เสพติดการโดดเรียน – Skip class addicted  •  Group 19  •  CP410844")
    set_run_font(footer_run, 8.5, False, "718096")

    document.core_properties.title = "รายงานโครงงานเกม 2D – เสพติดการโดดเรียน"
    document.core_properties.subject = "CP410844 Computer Game Development"
    document.core_properties.author = "Group 19"
    document.save(OUTPUT)
    print(f"REPORT_CREATED: {OUTPUT}")


if __name__ == "__main__":
    main()
